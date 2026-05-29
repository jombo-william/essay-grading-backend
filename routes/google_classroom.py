

"""
routes/google_classroom.py
Google Classroom Integration
"""
import json
import os
import secrets
import hashlib
import base64
from fastapi import APIRouter, Depends, HTTPException, Query
from pydantic import BaseModel
from fastapi.responses import RedirectResponse
from sqlalchemy.orm import Session
from auth_utils import require_teacher
from database import get_db
import models
import os
os.environ["OAUTHLIB_RELAX_TOKEN_SCOPE"] = "1"
_code_verifiers: dict = {}

router = APIRouter()

# ── Check if google packages are installed ────────────────────────────────────
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import Flow
    from googleapiclient.discovery import build
    GOOGLE_AVAILABLE = True
    print("✅ Google packages available")
except ImportError:
    GOOGLE_AVAILABLE = False
    print("❌ Google packages NOT installed — run: pip install google-auth google-auth-oauthlib google-api-python-client")

SCOPES = [
    "https://www.googleapis.com/auth/classroom.courses.readonly",
    "https://www.googleapis.com/auth/classroom.coursework.students",      # create/edit assignments
    "https://www.googleapis.com/auth/classroom.coursework.me",            # student submissions
    "https://www.googleapis.com/auth/classroom.student-submissions.students.readonly",
   #"https://www.googleapis.com/auth/classroom.grades",
    "https://www.googleapis.com/auth/classroom.rosters.readonly",
    "https://www.googleapis.com/auth/drive.readonly",
]

CLIENT_SECRETS_FILE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),  # routes/
    "..",                                          # go up one level to backend root
    "google_credentials.json"
)
print(f"📄 Creds path: {os.path.abspath(CLIENT_SECRETS_FILE)}")
print(f"📄 Size: {os.path.getsize(CLIENT_SECRETS_FILE)} bytes")
REDIRECT_URI = os.getenv(
    "GOOGLE_REDIRECT_URI",
    "http://localhost:8000/api/teacher/auth/google/callback"
)


# ── Text cleaner: nukes all NUL bytes and characters PostgreSQL rejects ───────
def clean_text(text: str) -> str:
    if not text:
        return ""
    # Remove null bytes explicitly (PostgreSQL rejects U+0000 in text columns)
    text = text.replace('\x00', '').replace('\u0000', '')
    # Remove other non-printable characters, keeping newlines/tabs
    text = ''.join(
        ch for ch in text
        if ch in '\n\r\t' or ord(ch) >= 32
    )
    # Round-trip through UTF-8 to strip lone surrogates and other invalid codepoints
    text = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
    return text.strip()


def extract_doc_bytes(content: bytes) -> str:
    """
    Extract readable text from a legacy .doc binary file.
    Tries python-docx first (works on many .doc files),
    then falls back to a latin-1 decode + printable-line filter.
    latin-1 is used because it maps bytes 1:1 and never produces NUL codepoints.
    """
    # Attempt 1 — python-docx (handles both .docx and some .doc files)
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if extracted.strip():
            return extracted
    except Exception:
        pass

    # Attempt 2 — latin-1 decode, keep only lines that look like real text
    try:
        text = content.decode("latin-1", errors="ignore")
        lines = [
            line for line in text.splitlines()
            if len(line.strip()) > 3
            and all(32 <= ord(c) < 127 or c in '\n\r\t' for c in line)
        ]
        return "\n".join(lines)
    except Exception:
        return ""


# ── Helper: load saved credentials for a teacher ─────────────────────────────
def get_credentials(teacher_id: int, db: Session):
    token_row = db.query(models.GoogleClassroomToken).filter_by(
        teacher_id=teacher_id
    ).first()

    if not token_row:
        raise HTTPException(
            status_code=401,
            detail="Google Classroom not connected. Please click 'Connect Google Classroom' first."
        )

    return Credentials(
        token         = token_row.access_token,
        refresh_token = token_row.refresh_token,
        token_uri     = token_row.token_uri,
        client_id     = token_row.client_id,
        client_secret = token_row.client_secret,
        scopes        = json.loads(token_row.scopes) if token_row.scopes else SCOPES,
    )

def get_gc_course_id_for_class(class_id: int, db: Session):
    """Get the linked Google Classroom course ID for a local class."""
    cls = db.query(models.Class).filter_by(id=class_id).first()
    return cls.gc_course_id if cls else None




# def sync_gc_students_to_db(course_id: str, service, db) -> dict:
#     """
#     Fetch all students from a Google Classroom course and create
#     local accounts for them if they don't exist yet.
#     Returns a dict of {gc_user_id: local_student_id}
#     """
#     import secrets
#     from passlib.context import CryptContext
#     pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

#     try:
#         result = service.courses().students().list(courseId=course_id).execute()
#         gc_students = result.get("students", [])
#         print(f"👥 Found {len(gc_students)} students in Google Classroom")
#     except Exception as e:
#         print(f"⚠️ Could not fetch students: {e}")
#         return {}

#     gc_id_to_local = {}

#     for gs in gc_students:
#         profile   = gs.get("profile", {})
#         gc_uid    = gs.get("userId", "")
#         name      = profile.get("name", {}).get("fullName", "Unknown Student")
#         email     = profile.get("emailAddress", f"{gc_uid}@classroom.google.com")

#         if not gc_uid:
#             continue

#         # Check if already linked
#         token_row = db.query(models.StudentGoogleToken).filter_by(
#             gc_user_id=gc_uid
#         ).first()

#         if token_row:
#             gc_id_to_local[gc_uid] = token_row.student_id
#             continue

#         # Check if user with this email already exists
#         existing_user = db.query(models.User).filter_by(email=email).first()

#         if existing_user:
#             student_id = existing_user.id
#         else:
#             # Create new student account
#             random_password = secrets.token_urlsafe(16)
#             hashed          = pwd_context.hash(random_password)
#             # new_user = models.User(
#             #     name     = name,
#             #     email    = email,
#             #     password = hashed,
#             #     role     = "student",
#             # )
#             from sqlalchemy import text

#             # Insert user via raw SQL to avoid PostgreSQL ENUM compile error
#             db.execute(
#                 text("INSERT INTO users (name, email, password, role) VALUES (:name, :email, :password, 'student')"),
#                 {"name": name, "email": email, "password": hashed}
#             )
#             db.flush()

#             # Get the newly created user
#             new_user = db.query(models.User).filter_by(email=email).first()
#             student_id = new_user.id
#             print(f"✅ Created student account: {name} ({email})")
                

#         if not existing_token:
#             db.add(models.StudentGoogleToken(
#                 student_id    = student_id,
#                 access_token  = "gc_sync",
#                 refresh_token = None,
#                 gc_user_id    = gc_uid,
#             ))

#         gc_id_to_local[gc_uid] = student_id

#     db.commit()
#     print(f"✅ Synced {len(gc_id_to_local)} students to local DB")
#     return gc_id_to_local


def sync_gc_students_to_db(course_id: str, service, db) -> dict:
    import secrets
    from passlib.context import CryptContext
    from sqlalchemy import text
    pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

    try:
        result = service.courses().students().list(courseId=course_id).execute()
        gc_students = result.get("students", [])
        print(f"👥 Found {len(gc_students)} students in Google Classroom")
    except Exception as e:
        print(f"⚠️ Could not fetch students: {e}")
        return {}

    gc_id_to_local = {}

    for gs in gc_students:
        profile = gs.get("profile", {})
        gc_uid  = gs.get("userId", "")
        name    = profile.get("name", {}).get("fullName", "Unknown Student")
        email   = profile.get("emailAddress", f"{gc_uid}@classroom.google.com")

        if not gc_uid:
            continue

        # Check if already linked
        token_row = db.query(models.StudentGoogleToken).filter_by(gc_user_id=gc_uid).first()
        if token_row:
            gc_id_to_local[gc_uid] = token_row.student_id
            continue

        # Check if user with this email already exists
        existing_user = db.query(models.User).filter_by(email=email).first()

        if existing_user:
            student_id = existing_user.id
        else:
            # Insert via raw SQL to bypass PostgreSQL ENUM compile error
            random_password = secrets.token_urlsafe(16)
            hashed = pwd_context.hash(random_password)
            db.execute(
                text("INSERT INTO users (name, email, password, role) VALUES (:name, :email, :password, 'student')"),
                {"name": name, "email": email, "password": hashed}
            )
            db.flush()
            new_user = db.query(models.User).filter_by(email=email).first()
            student_id = new_user.id
            print(f"✅ Created student account: {name} ({email})")

        # Link gc_user_id to local student
        existing_token = db.query(models.StudentGoogleToken).filter_by(student_id=student_id).first()
        if not existing_token:
            db.add(models.StudentGoogleToken(
                student_id    = student_id,
                access_token  = "gc_sync",
                refresh_token = None,
                gc_user_id    = gc_uid,
            ))

        gc_id_to_local[gc_uid] = student_id

    db.commit()
    print(f"✅ Synced {len(gc_id_to_local)} students to local DB")
    return gc_id_to_local




def create_gc_assignment(teacher_id: int, class_id: int, assignment, db: Session):
    """Create an assignment in Google Classroom and return the coursework ID."""
    gc_course_id = get_gc_course_id_for_class(class_id, db)
    if not gc_course_id:
        return None  # class not linked, skip silently

    try:
        creds   = get_credentials(teacher_id, db)
        service = build("classroom", "v1", credentials=creds)

        due = assignment.due_date
        coursework_body = {
            "title":       assignment.title,
            "description": assignment.description or assignment.instructions or "",
            "workType":    "ASSIGNMENT",
            "state":       "PUBLISHED",
            "maxPoints":   assignment.max_score,
            "dueDate": {
                "year":  due.year,
                "month": due.month,
                "day":   due.day,
            },
            "dueTime": {
                "hours":   due.hour,
                "minutes": due.minute,
                "seconds": 0,
                "nanos":   0,
            },
        }

        result = service.courses().courseWork().create(
            courseId=gc_course_id,
            body=coursework_body
        ).execute()

        return result.get("id")  # gc_coursework_id

    except Exception as e:
        print(f"⚠️ Could not create Google Classroom assignment: {e}")
        return None


    # ── ADD THIS FUNCTION to routes/google_classroom.py ──────────────────────────




def delete_gc_assignment(teacher_id: int, class_id: int, gc_coursework_id: str, db) -> bool:
    gc_course_id = get_gc_course_id_for_class(class_id, db)
    if not gc_course_id:
        return False

    try:
        creds   = get_credentials(teacher_id, db)
        service = build("classroom", "v1", credentials=creds)

        # Step 1 — check current state
        try:
            cw = service.courses().courseWork().get(
                courseId=gc_course_id,
                id=gc_coursework_id,
            ).execute()
        except Exception as e:
            if "404" in str(e):
                print(f"ℹ️ Coursework {gc_coursework_id} not found in Google Classroom — already deleted")
                return True  # not an error, just not there
            raise  # re-raise other errors

        # Step 2 — patch to DRAFT first if PUBLISHED
        if cw.get("state") != "DRAFT":
            service.courses().courseWork().patch(
                courseId=gc_course_id,
                id=gc_coursework_id,
                updateMask="state",
                body={"state": "DRAFT"},
            ).execute()
            print(f"📝 Set coursework {gc_coursework_id} to DRAFT")

        # Step 3 — now delete
        service.courses().courseWork().delete(
            courseId=gc_course_id,
            id=gc_coursework_id,
        ).execute()
        print(f"🗑️ Deleted coursework {gc_coursework_id} from Google Classroom")
        return True

    except Exception as e:
        print(f"⚠️ Could not delete Google Classroom coursework {gc_coursework_id}: {e}")
        return False


# ── GET /api/teacher/auth/google/classroom ────────────────────────────────────
@router.get("/auth/google/classroom")
def start_google_auth(ctx: dict = Depends(require_teacher), db: Session = Depends(get_db)):
    if not GOOGLE_AVAILABLE:
        raise HTTPException(status_code=500, detail="Google packages not installed.")

    if not os.path.exists(CLIENT_SECRETS_FILE):
        raise HTTPException(status_code=500, detail="google_credentials.json not found.")

    code_verifier = secrets.token_urlsafe(64)
    code_challenge = base64.urlsafe_b64encode(
        hashlib.sha256(code_verifier.encode()).digest()
    ).rstrip(b"=").decode()

    teacher_id = str(ctx["user"].id)

    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI
    )

    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent",
        state=teacher_id,
        code_challenge=code_challenge,
        code_challenge_method="S256",
    )

    existing = db.query(models.OAuthVerifier).filter_by(user_id=teacher_id).first()
    if existing:
        existing.code_verifier = code_verifier
    else:
        db.add(models.OAuthVerifier(user_id=teacher_id, code_verifier=code_verifier))
    db.commit()

    print(f"🔗 Google auth URL generated for teacher {teacher_id}")
    return {"auth_url": auth_url, "state": state}


# ── GET /api/teacher/auth/google/callback ─────────────────────────────────────
@router.get("/auth/google/callback")
def google_callback(
    code: str,
    state: str,
    db: Session = Depends(get_db)
):
    teacher_id = int(state)

    row = db.query(models.OAuthVerifier).filter_by(user_id=str(teacher_id)).first()
    code_verifier = row.code_verifier if row else None
    if row:
        db.delete(row)
        db.commit()

    if not code_verifier:
        raise HTTPException(status_code=400, detail="OAuth session expired or invalid. Please try connecting again.")

    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE,
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI
    )

    flow.fetch_token(code=code, code_verifier=code_verifier)
    creds = flow.credentials

    existing = db.query(models.GoogleClassroomToken).filter_by(
        teacher_id=teacher_id
    ).first()

    token_data = dict(
        access_token  = creds.token,
        refresh_token = creds.refresh_token,
        token_uri     = creds.token_uri,
        client_id     = creds.client_id,
        client_secret = creds.client_secret,
        scopes        = json.dumps(list(creds.scopes or SCOPES)),
    )

    if existing:
        for k, v in token_data.items():
            setattr(existing, k, v)
    else:
        db.add(models.GoogleClassroomToken(teacher_id=teacher_id, **token_data))

    db.commit()
    print(f"✅ Google tokens saved for teacher {teacher_id}")

    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
    return RedirectResponse(url=f"{frontend_url}?google_connected=true")

# ── GET /api/teacher/classroom/courses ───────────────────────────────────────
@router.get("/classroom/courses")
def get_courses(ctx: dict = Depends(require_teacher)):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    result = service.courses().list(
        teacherId    = "me",
        courseStates = ["ACTIVE"]
    ).execute()

    raw_courses = result.get("courses", [])
    print(f"📚 Found {len(raw_courses)} Google Classroom courses for teacher {user.id}")

    return {
        "success": True,
        "courses": [
            {
                "id":      c.get("id"),
                "name":    c.get("name"),
                "section": c.get("section", ""),
                "subject": c.get("descriptionHeading", ""),
            }
            for c in raw_courses
        ]
    }


# ── GET /api/teacher/classroom/courses/{course_id}/assignments ───────────────
@router.get("/classroom/courses/{course_id}/assignments")
def get_course_assignments(
    course_id: str,
    ctx: dict = Depends(require_teacher)
):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    result = service.courses().courseWork().list(
        courseId=course_id
    ).execute()

    work = result.get("courseWork", [])

    return {
        "success":     True,
        "assignments": [
            {
                "id":          a.get("id"),
                "title":       a.get("title"),
                "description": a.get("description", ""),
                "maxPoints":   a.get("maxPoints", 100),
            }
            for a in work
        ]
    }

# ── ADD HERE ──────────────────────────────────────────────────────────────────
@router.get("/classroom/debug-coursework/{course_id}/{cw_id}")
def debug_coursework(
    course_id: str,
    cw_id: str,
    ctx: dict = Depends(require_teacher)
):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    try:
        result = service.courses().courseWork().get(
            courseId=course_id,
            id=cw_id,
        ).execute()
        return {"found": True, "data": result}
    except Exception as e:
        return {"found": False, "error": str(e)}


# ── POST /api/teacher/classroom/courses/{course_id}/sync ─────────────────────
@router.post("/classroom/courses/{course_id}/sync")
def sync_gc_assignments(
    course_id: str,
    ctx: dict = Depends(require_teacher)
):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    # Find the local class linked to this gc_course_id
    cls = db.query(models.Class).filter_by(gc_course_id=course_id).first()
    if not cls:
        raise HTTPException(
            status_code=404,
            detail="No local class is linked to this Google Classroom course. Please link it first."
        )

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    # Fetch all coursework from Google Classroom
    result = service.courses().courseWork().list(
        courseId=course_id
    ).execute()

    gc_assignments = result.get("courseWork", [])
    print(f"📥 Found {len(gc_assignments)} assignments in Google Classroom")

    created = 0
    skipped = 0

    for gca in gc_assignments:
        gc_cw_id = gca.get("id")

        # Skip if already in our DB
        existing = db.query(models.Assignment).filter_by(
            gc_coursework_id=gc_cw_id
        ).first()
        if existing:
            skipped += 1
            continue

        # Parse due date
        due_date = None
        if "dueDate" in gca and "dueTime" in gca:
            d = gca["dueDate"]
            t = gca["dueTime"]
            try:
                from datetime import datetime
                from zoneinfo import ZoneInfo
                due_date = datetime(
                    d.get("year",  2025),
                    d.get("month", 1),
                    d.get("day",   1),
                    t.get("hours",   0),
                    t.get("minutes", 0),
                    tzinfo=ZoneInfo("Africa/Blantyre")
                )
            except Exception:
                from datetime import datetime, timezone, timedelta
                due_date = datetime.now(timezone.utc) + timedelta(days=7)
        else:
            from datetime import datetime, timezone, timedelta
            due_date = datetime.now(timezone.utc) + timedelta(days=7)

        # Create assignment in local DB
        new_assignment = models.Assignment(
            teacher_id         = user.id,
            class_id           = cls.id,
            title              = gca.get("title", "Untitled"),
            description        = gca.get("description", ""),
            instructions       = gca.get("description", "Imported from Google Classroom"),
            reference_material = None,
            max_score          = int(gca.get("maxPoints", 100)),
            due_date           = due_date,
            rubric             = None,
            gc_coursework_id   = gc_cw_id,
        )
        db.add(new_assignment)
        created += 1

    db.commit()




    # ── Also sync students and their submissions ──────────────────────────
    gc_id_to_local = sync_gc_students_to_db(course_id, service, db)

    # Enroll synced students into the local class
    for gc_uid, student_id in gc_id_to_local.items():
        exists = db.query(models.ClassEnrollment).filter_by(
            class_id=cls.id, student_id=student_id
        ).first()
        if not exists:
            db.add(models.ClassEnrollment(class_id=cls.id, student_id=student_id))
    db.commit()

    # Fetch all submissions for all assignments in this course
    local_assignments = db.query(models.Assignment).filter_by(
        class_id=cls.id
    ).all()

    submissions_synced = 0
    for local_assignment in local_assignments:
        if not local_assignment.gc_coursework_id:
            continue
        try:
            subs_result = service.courses().courseWork().studentSubmissions().list(
                courseId     = course_id,
                courseWorkId = local_assignment.gc_coursework_id,
                states       = ["TURNED_IN", "RETURNED"]
            ).execute()
            gc_subs = subs_result.get("studentSubmissions", [])

            for gc_sub in gc_subs:
                gc_uid    = gc_sub.get("userId", "")
                student_id = gc_id_to_local.get(gc_uid)
                if not student_id:
                    continue

                # Check if submission already exists
                existing_sub = db.query(models.Submission).filter_by(
                    assignment_id = local_assignment.id,
                    student_id    = student_id,
                ).first()

                if existing_sub:
                    continue  # already synced

                # Get assigned grade if any
                assigned_grade = gc_sub.get("assignedGrade")
                draft_grade    = gc_sub.get("draftGrade")
                gc_score       = assigned_grade or draft_grade

                new_sub = models.Submission(
                    assignment_id      = local_assignment.id,
                    student_id         = student_id,
                    essay_text         = "[Submitted via Google Classroom — use Grade button to extract and grade]",
                    submit_mode        = "upload",
                    file_name          = f"gc_{gc_uid}",
                    ai_score           = int(gc_score) if gc_score is not None else None,
                    ai_feedback        = None,
                    ai_detection_score = 0,
                    status             = "ai_graded" if gc_score is not None else "submitted",
                )
                db.add(new_sub)
                submissions_synced += 1

        except Exception as e:
            print(f"⚠️ Could not sync submissions for assignment {local_assignment.id}: {e}")

    db.commit()
    print(f"✅ Synced {submissions_synced} submissions from Google Classroom")
    # ─────────────────────────────────────────────────────────────────────

    print(f"✅ Sync complete: {created} created, {skipped} already existed")
    return {
        "success": True,
        "created": created,
        "skipped": skipped,
        "students_synced": len(gc_id_to_local),
        "submissions_synced": submissions_synced,
        "message": f"{created} assignment(s) imported, {len(gc_id_to_local)} students synced, {submissions_synced} submissions synced.",
    }


    print(f"✅ Sync complete: {created} created, {skipped} already existed")
    return {
        "success": True,
        "created": created,
        "skipped": skipped,
        "message": f"{created} assignment(s) imported, {skipped} already existed.",
    }        

# ── POST /api/teacher/classroom/courses/{course_id}/assignments/{cw_id}/grade ─
@router.post("/classroom/courses/{course_id}/assignments/{coursework_id}/grade")
def import_and_grade(
    course_id:           str,
    coursework_id:       str,
    local_assignment_id: int = Query(...),
    ctx: dict = Depends(require_teacher)
):
    from routes.ai_grader import grade_with_ai

    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    assignment = db.query(models.Assignment).filter(
        models.Assignment.id         == local_assignment_id,
        models.Assignment.teacher_id == user.id,
    ).first()

    if not assignment:
        raise HTTPException(status_code=404, detail="Local assignment not found")

    creds         = get_credentials(user.id, db)
    classroom_svc = build("classroom", "v1", credentials=creds)
    drive_svc     = build("drive",     "v3", credentials=creds)

    subs_result = classroom_svc.courses().courseWork().studentSubmissions().list(
        courseId     = course_id,
        courseWorkId = coursework_id,
        states       = ["TURNED_IN"]
    ).execute()

    student_subs = subs_result.get("studentSubmissions", [])
    print(f"📥 Found {len(student_subs)} submissions in Google Classroom")

    results = []

    # Track which local student_ids were graded via GC
    # so we don't re-grade them in the local pass
    gc_graded_student_ids = set()

    # ── Grade Google Classroom submissions ────────────────────────────────────
   

    for gs in student_subs:
        gc_uid      = gs.get("userId", "unknown")
        essay_text  = ""

        # Look up student name from local DB
        gc_token_row = db.query(models.StudentGoogleToken).filter_by(gc_user_id=gc_uid).first()
        if gc_token_row:
            student_user_row = db.query(models.User).filter_by(id=gc_token_row.student_id).first()
            student_name = student_user_row.name if student_user_row else gc_uid
        else:
            student_name = gc_uid
        attachments = gs.get("assignmentSubmission", {}).get("attachments", [])

        # ── Resolve real student name ─────────────────────────────────────
        gc_token = db.query(models.StudentGoogleToken).filter_by(
            gc_user_id=gc_uid
        ).first()
        actual_student_id = gc_token.student_id if gc_token else None

        student_name = "Unknown"
        if actual_student_id:
            local_user = db.query(models.User).filter_by(id=actual_student_id).first()
            if local_user:
                student_name = local_user.name

        for att in attachments:
            if "driveFile" in att:
                file_id = att["driveFile"]["id"]
                try:
                    file_meta = drive_svc.files().get(
                        fileId = file_id,
                        fields = "mimeType, name"
                    ).execute()
                    mime = file_meta.get("mimeType", "")
                    file_name = file_meta.get("name", "")

                    print(f"📎 Processing file: {file_name} (mime: {mime})")

                    if mime == "application/vnd.google-apps.document":
                        # Google Doc — export as plain text (always clean)
                        content = drive_svc.files().export(
                            fileId=file_id, mimeType="text/plain"
                        ).execute()
                        essay_text += content.decode("utf-8", errors="ignore")
                    

                    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        # .docx file — use python-docx for clean extraction
                        content = drive_svc.files().get_media(fileId=file_id).execute()
                        try:
                            import io
                            from docx import Document
                            doc = Document(io.BytesIO(content))
                            extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                            essay_text += extracted
                            print(f"✅ Extracted .docx text ({len(extracted)} chars)")
                        except Exception as docx_err:
                            print(f"⚠️ python-docx failed for {file_name}: {docx_err} — falling back")
                            essay_text += content.decode("utf-8", errors="ignore")

                    elif mime in (
                        "application/msword",
                        "application/vnd.ms-word",
                        "application/x-msword",
                    ):
                        # Legacy .doc file — binary format, must be handled carefully
                        content = drive_svc.files().get_media(fileId=file_id).execute()
                        extracted = extract_doc_bytes(content)
                        essay_text += extracted
                        print(f"✅ Extracted legacy .doc text ({len(extracted)} chars)")

                    elif mime == "application/pdf":
                        content = drive_svc.files().get_media(fileId=file_id).execute()
                        try:
                            import io
                            import pypdf
                            reader = pypdf.PdfReader(io.BytesIO(content))
                            extracted = ""
                            for page in reader.pages:
                                extracted += page.extract_text() or ""
                            extracted = extracted.strip()
                            if len(extracted) > 50:
                                print(f"📄 PDF extracted {len(extracted)} chars: {extracted[:100]}")
                                essay_text += extracted
                            else:
                                print(f"⚠️ pypdf got nothing — falling back to raw decode")
                                essay_text += content.decode("utf-8", errors="ignore")
                        except Exception as pdf_err:
                            print(f"⚠️ pypdf failed: {pdf_err}")
                            essay_text += content.decode("utf-8", errors="ignore")

                    elif "text" in mime:
                        content = drive_svc.files().get_media(fileId=file_id).execute()
                        essay_text += content.decode("utf-8", errors="ignore")
                   
                   
                    elif mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
                        # Word document — extract with python-docx
                        try:
                            import io
                            from docx import Document
                            content = drive_svc.files().get_media(fileId=file_id).execute()
                            doc = Document(io.BytesIO(content))
                            essay_text += "\n".join([para.text for para in doc.paragraphs])
                            print(f"📄 Word doc extracted {len(essay_text)} chars")
                        except Exception as docx_err:
                            print(f"⚠️ Word doc extraction failed: {docx_err}")                   

                    else:
                        # Unknown type — attempt utf-8 decode as last resort
                        content = drive_svc.files().get_media(fileId=file_id).execute()
                        essay_text += content.decode("utf-8", errors="ignore")

                    print(f"✅ Read file {file_id} (type: {mime})")

                except Exception as e:
                    print(f"⚠️ Could not read Drive file {file_id}: {e}")

        # ── Clean extracted text before any DB interaction ────────────────
        # This MUST happen before the essay_text.strip() check and before
        # any DB save. clean_text removes NUL bytes that cause PostgreSQL
        # "string literal cannot contain NUL" errors.
        essay_text = clean_text(essay_text)

        if not essay_text.strip():
            results.append({
                "student_name":      student_name,
                "google_student_id": gc_uid,
                "error":             "No text content found in submission",
                "status":            "skipped",
                "source":            "google_classroom",
            })
            continue

        try:
            word_count = len(essay_text.split())
            from routes.grading_prompt import build_grading_prompt
            prompt = build_grading_prompt(assignment, essay_text, word_count)
            grade  = grade_with_ai(
                prompt=prompt,
                assignment=assignment,
                essay_text=essay_text,
                word_count=word_count,
            )

            # ── Save to DB — always clean text slice before saving ────────
            safe_essay = clean_text(essay_text)[:5000]

            existing_sub = db.query(models.Submission).filter(
                models.Submission.assignment_id == assignment.id,
                models.Submission.file_name     == f"gc_{gc_uid}",
            ).first()

            if existing_sub:
                existing_sub.essay_text         = safe_essay
                existing_sub.ai_score           = grade["score"]
                existing_sub.ai_feedback        = grade["feedback"]
                existing_sub.ai_detection_score = 0
                existing_sub.status             = "ai_graded"
                db.commit()
                if existing_sub.student_id:
                    gc_graded_student_ids.add(existing_sub.student_id)

            elif actual_student_id:
                gc_graded_student_ids.add(actual_student_id)
                check = db.query(models.Submission).filter_by(
                    assignment_id = assignment.id,
                    student_id    = actual_student_id,
                ).first()
                if check:
                    check.essay_text         = safe_essay
                    check.ai_score           = grade["score"]
                    check.ai_feedback        = grade["feedback"]
                    check.ai_detection_score = 0
                    check.status             = "ai_graded"
                    check.file_name          = f"gc_{gc_uid}"
                else:
                    db.add(models.Submission(
                        assignment_id      = assignment.id,
                        student_id         = actual_student_id,
                        essay_text         = safe_essay,
                        submit_mode        = "upload",
                        file_name          = f"gc_{gc_uid}",
                        ai_score           = grade["score"],
                        ai_feedback        = grade["feedback"],
                        ai_detection_score = 0,
                        status             = "ai_graded",
                    ))
                db.commit()
            else:
                print(f"⚠️ Could not find local student for GC user {gc_uid} — skipping DB save")

            # ── Push grade back to Google Classroom ───────────────────────
            try:
                gc_sub_list = classroom_svc.courses().courseWork().studentSubmissions().list(
                    courseId     = course_id,
                    courseWorkId = coursework_id,
                    userId       = gc_uid,
                ).execute()
                gc_subs = gc_sub_list.get("studentSubmissions", [])
                if gc_subs:
                    gc_sub_id = gc_subs[0]["id"]
                    classroom_svc.courses().courseWork().studentSubmissions().patch(
                        courseId     = course_id,
                        courseWorkId = coursework_id,
                        id           = gc_sub_id,
                        updateMask   = "assignedGrade,draftGrade",
                        body={
                            "assignedGrade": grade["score"],
                            "draftGrade":    grade["score"],
                        },
                    ).execute()
                    print(f"✅ Grade {grade['score']} posted to Google Classroom for {gc_uid}")
            except Exception as grade_err:
                print(f"⚠️ Could not post grade to Google Classroom: {grade_err}")

            results.append({
                "student_name":      student_name,
                "google_student_id": gc_uid,
                "score":             grade["score"],
                "feedback":          grade["feedback"],
                "status":            "graded",
                "source":            "google_classroom",
            })
            print(f"✅ Graded GC submission for {student_name} → {grade['score']}/{assignment.max_score}")

        except Exception as e:
            db.rollback()
            print(f"❌ Grading failed for {student_name}: {e}")
            results.append({
                "student_name":      student_name,
                "google_student_id": gc_uid,
                "error":             str(e),
                "status":            "failed",
                "source":            "google_classroom",
            })

    # ── Grade local submissions — SKIP anyone already graded via GC ──────────
    local_subs = db.query(models.Submission, models.User).join(
        models.User, models.User.id == models.Submission.student_id
    ).filter(
        models.Submission.assignment_id == assignment.id
    ).all()

    print(f"📥 Found {len(local_subs)} submissions in local DB")

    # for sub, student_user in local_subs:
    #     essay_text = sub.essay_text
    #     if not essay_text or not essay_text.strip():
    #         continue

    for sub, student_user in local_subs:

        # ── Skip if this student was already graded via Google Classroom ──
        if student_user.id in gc_graded_student_ids:
            print(f"⏭️ Skipping {student_user.name} — already graded via Google Classroom")
            continue

        essay_text = sub.essay_text
        if not essay_text or not essay_text.strip():
            continue

        # Skip placeholder text from sync
        if essay_text.strip().startswith("[Submitted via Google Classroom"):
            continue

        try:
            word_count = len(essay_text.split())
            from routes.grading_prompt import build_grading_prompt
            prompt = build_grading_prompt(assignment, essay_text, word_count)
            grade  = grade_with_ai(
                prompt=prompt,
                assignment=assignment,
                essay_text=essay_text,
                word_count=word_count,
            )

            sub.ai_score           = grade["score"]
            sub.ai_feedback        = grade["feedback"]
            sub.ai_detection_score = 0
            sub.status             = "ai_graded"
            db.commit()

            results.append({
                "student_name":      student_user.name,
                "google_student_id": f"local_{student_user.id}",
                "score":             grade["score"],
                "feedback":          grade["feedback"],
                "status":            "graded",
                "source":            "local",
            })
            print(f"✅ Graded local submission for {student_user.name} → {grade['score']}/{assignment.max_score}")

        except Exception as e:
            db.rollback()
            print(f"❌ Local grading failed for {student_user.name}: {e}")
            results.append({
                "student_name":      student_user.name,
                "google_student_id": f"local_{student_user.id}",
                "error":             str(e),
                "status":            "failed",
                "source":            "local",
            })

    return {
        "success":      True,
        "total_graded": len([r for r in results if r["status"] == "graded"]),
        "results":      results,
    }

# ── GET /api/teacher/classroom/status ────────────────────────────────────────
@router.get("/classroom/status")
def check_connection_status(ctx: dict = Depends(require_teacher)):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    token_row = db.query(models.GoogleClassroomToken).filter_by(
        teacher_id=user.id
    ).first()

    return {
        "connected": token_row is not None,
        "message":   "Google Classroom connected" if token_row else "Not connected"
    }


# ── POST /api/teacher/classes/{class_id}/link-google ─────────────────────────

class LinkGoogleRequest(BaseModel):
    gc_course_id: str


@router.post("/classes/{class_id}/link-google")
def link_google_course(
    class_id: int,
    body: LinkGoogleRequest,
    ctx: dict = Depends(require_teacher),
):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    cls = db.query(models.Class).filter_by(id=class_id).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found.")

    cls.gc_course_id = body.gc_course_id
    db.commit()

    return {"success": True, "message": "Google Classroom course linked to class."}


# ── POST /api/teacher/classroom/courses/{course_id}/enroll-students ──────────
@router.post("/classroom/courses/{course_id}/enroll-students")
def enroll_gc_students(
    course_id:      str,
    local_class_id: int = Query(...),
    ctx: dict = Depends(require_teacher)
):
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    result = service.courses().students().list(
        courseId=course_id
    ).execute()

    gc_students = result.get("students", [])
    print(f"👥 Found {len(gc_students)} students in GC course {course_id}")

    enrolled = 0
    created  = 0
    skipped  = 0

    for gc_student in gc_students:
        profile   = gc_student.get("profile", {})
        gc_uid    = gc_student.get("userId")
        email     = profile.get("emailAddress", "")
        full_name = profile.get("name", {}).get("fullName", f"Student {gc_uid}")

        if not email:
            print(f"⚠️ No email for GC user {gc_uid} — skipping")
            skipped += 1
            continue

        # Step 1 — find or create local user
        local_user = db.query(models.User).filter_by(email=email).first()

        if not local_user:
            import bcrypt
            random_pw = secrets.token_urlsafe(12)
            hashed    = bcrypt.hashpw(random_pw.encode(), bcrypt.gensalt()).decode()

            local_user = models.User(
                name     = full_name,
                email    = email,
                password = hashed,
                role     = "student",
            )
            db.add(local_user)
            db.flush()
            created += 1
            print(f"✅ Created local account for {full_name} ({email})")

        # Step 2 — link gc_user_id for grading matching
        gc_token = db.query(models.StudentGoogleToken).filter_by(
            student_id=local_user.id
        ).first()

        if not gc_token:
            db.add(models.StudentGoogleToken(
                student_id = local_user.id,
                gc_user_id = gc_uid,
            ))

        # Step 3 — enroll in local class if not already enrolled
        already_enrolled = db.query(models.ClassEnrollment).filter_by(
            class_id   = local_class_id,
            student_id = local_user.id,
        ).first()

        if not already_enrolled:
            db.add(models.ClassEnrollment(
                class_id   = local_class_id,
                student_id = local_user.id,
            ))
            enrolled += 1

    db.commit()

    print(f"✅ GC sync done: {created} created, {enrolled} enrolled, {skipped} skipped")

    return {
        "success":  True,
        "created":  created,
        "enrolled": enrolled,
        "skipped":  skipped,
        "message":  f"{created} accounts created, {enrolled} enrolled, {skipped} skipped",
    }


# ── POST /api/teacher/classroom/courses/{course_id}/sync ─────────────────────
@router.post("/classroom/courses/{course_id}/sync")
def sync_gc_course(
    course_id:      str,
    local_class_id: int = Query(...),
    ctx: dict = Depends(require_teacher)
):
    """
    Sync a Google Classroom course with a local class:
    1. Enroll all GC students locally
    2. Import all GC assignments locally
    """
    user: models.User = ctx["user"]
    db:   Session     = ctx["db"]

    creds   = get_credentials(user.id, db)
    service = build("classroom", "v1", credentials=creds)

    # ── Sync students ─────────────────────────────────────────────────────────
    students_result = service.courses().students().list(
        courseId=course_id
    ).execute()

    gc_students = students_result.get("students", [])
    students_created  = 0
    students_enrolled = 0

    for gc_student in gc_students:
        profile   = gc_student.get("profile", {})
        gc_uid    = gc_student.get("userId")
        email     = profile.get("emailAddress", "")
        full_name = profile.get("name", {}).get("fullName", f"Student {gc_uid}")

        if not email:
            continue

        local_user = db.query(models.User).filter_by(email=email).first()
        if not local_user:
            import bcrypt
            hashed = bcrypt.hashpw(
                secrets.token_urlsafe(12).encode(), bcrypt.gensalt()
            ).decode()
            local_user = models.User(
                name=full_name, email=email,
                password=hashed, role="student",
            )
            db.add(local_user)
            db.flush()
            students_created += 1

        # Link gc_user_id
        gc_token = db.query(models.StudentGoogleToken).filter_by(
            student_id=local_user.id
        ).first()
        if not gc_token:
            db.add(models.StudentGoogleToken(
                student_id=local_user.id,
                gc_user_id=gc_uid,
            ))

        # Enroll in local class
        already = db.query(models.ClassEnrollment).filter_by(
            class_id=local_class_id, student_id=local_user.id
        ).first()
        if not already:
            db.add(models.ClassEnrollment(
                class_id=local_class_id, student_id=local_user.id
            ))
            students_enrolled += 1

    # ── Sync assignments ──────────────────────────────────────────────────────
    work_result = service.courses().courseWork().list(
        courseId=course_id
    ).execute()

    gc_assignments     = work_result.get("courseWork", [])
    assignments_synced = 0

    for gca in gc_assignments:
        gc_cw_id = gca.get("id")

        # Skip if already imported
        exists = db.query(models.Assignment).filter_by(
            gc_coursework_id=gc_cw_id
        ).first()
        if exists:
            continue

        due_date = None
        if gca.get("dueDate"):
            d = gca["dueDate"]
            t = gca.get("dueTime", {})
            from datetime import datetime
            due_date = datetime(
                d.get("year", 2025), d.get("month", 1), d.get("day", 1),
                t.get("hours", 23), t.get("minutes", 59), 0,
            )

        new_assignment = models.Assignment(
            teacher_id       = user.id,
            class_id         = local_class_id,
            title            = gca.get("title", "Untitled"),
            description      = gca.get("description", ""),
            instructions     = gca.get("description", "Imported from Google Classroom"),
            max_score        = int(gca.get("maxPoints", 100)),
            due_date         = due_date,
            gc_coursework_id = gc_cw_id,
            is_active        = True,
        )
        db.add(new_assignment)
        assignments_synced += 1

    db.commit()

    print(f"✅ Sync done: {students_created} students created, {students_enrolled} enrolled, {assignments_synced} assignments imported")

    return {
        "success":            True,
        "students_created":   students_created,
        "students_enrolled":  students_enrolled,
        "assignments_synced": assignments_synced,
        "message":            f"Sync complete — {students_enrolled} students enrolled, {assignments_synced} assignments imported",
    }
